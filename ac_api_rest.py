import os
import io
import sys
import ast
import uuid
import json
import grpc
import time
import base64
import logging
import requests
import docx2txt 
import numpy as np
import pickle as pl
from pracs import map_func
#from pdfcode import *
from docx import Document
from flask_cors import CORS
from datetime import datetime
from datetime import timedelta
from pdfreader import SimplePDFViewer
from flask_httpauth import HTTPBasicAuth
from flask import Flask, jsonify, abort, request, make_response, Response
from werkzeug.exceptions import BadRequest
from filelock import FileLock 
from ac_base import set_dl_param_code, get_dl_param_code, del_dl_param_code
from ac_base import read_local_dl_params, extract_license
from dotenv import load_dotenv 
from pydantic_settings import BaseSettings
from typing import Optional

app = Flask(__name__)

# Load environment variables from the .env files before using them
load_dotenv(dotenv_path= "/home/ac/ai_ac_service/api/init_param.env")  # Load the first .env file
load_dotenv(dotenv_path="/home/ac/ai_ac_service/api/init_dl_param.env")  # Load the second .env file

# Function to print environment variables after loading
def print_environment_variables():
    print("\nLoaded Environment Variables:")
    for key, value in os.environ.items():
        print(f"{key}: {value}")
        
# Print the environment variables after loading them
print_environment_variables()

# Define a Pydantic model for reading general initialization parameters
class InitParams(BaseSettings):
    ai_type: str
    ai_type_dl: str
    ai_type_knn: str
    db_type: str

# Define a Pydantic model for reading deep learning related initialization parameters
class DLInitParams(BaseSettings):
    description: str
    reject_thres: float
    max_num_results: int
    oc_reject_thres: float
    oc_num_single_results: int

# Load the environment variables using Pydantic models
params = InitParams()
dl_params = DLInitParams()

# function/method for checking if this code is running within docker
def is_docker():
    path = '/proc/self/cgroup'
    return (
        os.path.exists('/.dockerenv') or
        os.path.isfile(path) and any('docker' in line for line in open(path))
    )

# Rest of your code continues...


if is_docker():
    sys.path.insert(0, '../basicdata/')
else:
    # get current directory of AI service
    curr_dir = os.getcwd()
    print('Current AI service directory: ', curr_dir)
    # get basic directory
    basic_dir = curr_dir+'/basicdata/'
    print('Basic directory: ', basic_dir)
    sys.path.insert(0, basic_dir)

def token_func(ts_model_name):
    if is_docker():
        tokenize_path = '../models/'+ts_model_name+'/tokenizer.json'
    else:
        tokenize_path = './tfserving/tfservingmodels/'+ts_model_name+'/tokenizer.json'
    with open(tokenize_path) as f:
        data = json.load(f)
    return data 

 
 
# Flask route for testing the loaded parameters
@app.route('/test_params', methods=['GET'])
def test_params():
    response = {
        "description": params.description,
        "ai_type": params.ai_type,
        "ai_type_dl": params.ai_type_dl,
        "ai_type_knn": params.ai_type_knn,
        "db_type": params.db_type,
        "reject_thres": dl_params.reject_thres,
        "max_num_results": dl_params.max_num_results,
        "oc_reject_thres": dl_params.oc_reject_thres,
        "oc_num_single_results": dl_params.oc_num_single_results
    }
    return jsonify(response)
 
# Original 1126 lines code below (commented out)
 
from basic_functions_cython import set_module_dir
from basic_functions_cython import write_log_api
from basic_functions_cython import write_log_api_file
from basic_functions_cython import init_global_file_data
# from basic_functions_cython import read_init_param  
from basic_functions_cython import write_dl_init_param
from basic_functions_cython import set_dl_init_info_file
from basic_functions_cython import get_dl_init_info_file
# from basic_functions_cython import read_dl_init_param
from basic_functions_cython import delete_dl_info_data_file
from basic_functions_cython import TensorServingGRPC_Client
from basic_functions_cython import set_global_file_data
from basic_db_functions_cython import db_session
from basic_db_functions_cython import get_init_status
from basic_db_functions_cython import set_init_status
from basic_db_functions_cython import set_dl_init_data
from basic_db_functions_cython import set_version_info
from basic_db_functions_cython import set_dl_init_info
from basic_db_functions_cython import get_dl_init_data
 
def function_write_file_extract_data(binary_content, suffix):
    unique_time_stamp = str(int(time.time()))
    filename = unique_time_stamp+suffix
    data_file = open(filename, "wb")
    data_file.write(binary_content)
    data_file.close()
    if os.path.isfile(filename):
        data_file = open((filename), "r")
        file = data_file.name
        if file.endswith('.docx'):
            extracted_string_data = docx2txt.process(data_file.name)
        elif file.endswith('.pdf'):
            logging_text = 'Please do not use PDF file'
            if db_type == "File":
                write_log_api_file(module_dir, module, logging_text)  
            else:
                write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
            #extracted_string_data = Convert.convert_pdf_to_txt(data_file.name)
            extracted_string_data = 'No Data'
        elif file.endswith('.txt'):
            extracted_string_data = data_file.read()
        else:
            extracted_string_data = data_file.read()
    else:
        print('file not found')
    os.remove(filename)
    new_string = extracted_string_data
    chars_to_remove = [ "@", "1", "2", "3", "4", "5", "6", "7", "8", "9", "0", ".", ",", ";", ":", "-", "\r", "\n", ">", "<"]
    for element in new_string:
        if element in chars_to_remove:
            new_string = new_string.replace(element, ' ')
    special_chars_to_remove = ["&"]
    for element in new_string:
        if element in special_chars_to_remove:
            index_special_char = int(new_string.index('&'))
            if index_special_char==0:
                new_string = ' ' + new_string[3 : : ]
            elif index_special_char==len(new_string)-2:
                new_string = new_string.replace(element, ' ')
            elif index_special_char==len(new_string)-1:
                new_string = new_string.replace(element, ' ')
            else:
                start = index_special_char
                stop = index_special_char+2
                new_string = new_string[0: start:] + ' ' + new_string[stop + 1::]
    extracted_string_data = new_string
    return extracted_string_data
 
def function_extract_data_from_byte_stream(binary_content, suffix):
    extracted_string_data=''
    if suffix=='pdf':
        viewer = SimplePDFViewer(binary_content)
        viewer.render()
        string_data = viewer.canvas.strings
        string_data = ''.join(string_data)
        extracted_string_data = string_data
    elif suffix=='docx':
        string_data = ''
        doc = Document(io.BytesIO(binary_content))
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        string_data = ''.join(full_text)
        extracted_string_data = string_data
    else:
        string_data = binary_content.decode('utf-8')
        extracted_string_data = string_data
    new_string = extracted_string_data
    chars_to_remove = [ "@", "1", "2", "3", "4", "5", "6", "7", "8", "9", "0", ".", ",", ";", ":", "-", "\r", "\n", ">", "<"]
    for element in new_string:
        if element in chars_to_remove:
            new_string = new_string.replace(element, ' ')
    special_chars_to_remove = ["&"]
    for element in new_string:
        if element in special_chars_to_remove:
            index_special_char = int(new_string.index('&'))
            if index_special_char==0:
                new_string = ' ' + new_string[3 : : ]
            elif index_special_char==len(new_string)-2:
                new_string = new_string.replace(element, ' ')
            elif index_special_char==len(new_string)-1:
                new_string = new_string.replace(element, ' ')
            else:
                start = index_special_char
                stop = index_special_char+2
                new_string = new_string[0: start:] + ' ' + new_string[stop + 1::]
    extracted_string_data = new_string
    return extracted_string_data
 
def dl_get_text_result_body(module, text_content, ip_address_tf, port_tf, dl_init_ref, user_id, dl_param_id=''):
   
    module = module + '/dl_get_text_result_body '
   
    start_timer_1 = time.time()
 
    if db_type == 'File':
        status_current = 'OK'
    else:
        status_data = get_init_status(session, StatusData, status_datas_schema, user_id)
        if len(status_data) > 0:
            status_current = status_data['status']
            status_info_current = status_data['status_info']
            dl_init_status_current = status_data['dl_init_status']
            dl_init_status_info_current = status_data['dl_init_status_info']
        else:
            dl_results = []
            dl_result = {'id': 0, 'dl_result':'-1000', 'dl_qval':'-1000', 'dl_info':'DL init status not available'}
            dl_results.append(dl_result)
            return jsonify(dl_results)
 
        if dl_init_status_current == 'NOK':
            dl_results = []
            dl_result = {'id': 0, 'dl_result':'-1000', 'dl_qval':-1000, 'dl_info':'DL classifier not initialized'}
            dl_results.append(dl_result)
            return jsonify(dl_results)
 
    if status_current == 'OK':
        if db_type == "File":
            GetDlInitDataInfoOK, dl_init_data_info_list = get_dl_init_info_file(module_dir, dl_init_ref)
            if GetDlInitDataInfoOK:
                ts_model_name = dl_init_data_info_list[0].get('dlinit_tsmodelname')
                ts_model_version = dl_init_data_info_list[0].get('dlinit_tsmodelversion')
                ts_model_classes = dl_init_data_info_list[0].get('dlinit_tsmodelclasses')
                logging_text = ' TF Serving Model, Version & Classes successfully loaded from data/dl_init_info_xxxxxxxx.json file'
                write_log_api_file(module_dir, module, logging_text)
            else:
                logging_text = ' TF Serving Model, Version & Classes could not be successfully loaded from data/dl_init_info_xxxxxxxx.json file'
                write_log_api_file(module_dir, module, logging_text)
                ts_model_name = ''
                ts_model_version = ''
                ts_model_classes = ''
        else:
            ts_model_name, ts_model_version, ts_model_classes = get_dl_init_data(session, DlInit, dl_inits_schema)
    else:
        dl_results = []
        dl_result = {'id': 0, 'dl_result':'-1000', 'dl_qval':'-1000', 'dl_info':'DL not initialized'}
        dl_results.append(dl_result)
        return jsonify(dl_results)
   
    start_timer_2 = time.time()
    headers = {'Content-Type': 'application/json'}
    tokens= token_func(ts_model_name)
    string_data = {"dirAddress":text_content, "tokenizer": tokens}
 
    #preparing URL for connecting to ai_ac_datapreprocess_api_rest container
    if is_docker():
        ## this works for DOCKER server
        ip_dp='ai_datapreprocess_api_rest'
    else:
        ## this work for MANUAL server
        ip_dp='127.0.0.1'
 
    dp_url='https://'+ip_dp+':6355/ac/api/data_preprocessing'
   
    # following line is for insecure communication
    #classify= requests.post(dp_url.replace('ps:','p:'), headers=headers, json=string_data, auth=(server_username, server_password))
   
    # following line is for secure communication
    classify= requests.post(dp_url, headers=headers, json=string_data, auth=(server_username, server_password), cert=('server_local.crt', 'server_local.key'), verify=False)
     
    end_timer_2 = time.time()
    timer_2 = end_timer_2 - start_timer_2
    logging_text = ' Time for data preprocessing from another service: ' + str(timer_2) + ' seconds'
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)  
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
   
    x_test = np.load('numpy/x_test.npy')
   
    # read_init => need  to remove and read from dotenv file
    #read global DL init data
    # reject_thres, max_num_results, oc_reject_thres, oc_num_single_results = read_dl_init_param(module_dir, db_type, session, StatusData, status_datas_schema)
    logging_text = ' DL reject threshold: '+str(reject_thres)
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)  
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
    logging_text = ' DL maximum number of results per decision: '+str(max_num_results)
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)  
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
 
    # read local DL parameters
    if dl_param_id != '0':
        reject_thres,max_num_results,oc_reject_thres,oc_num_single_results,error_dl_params = read_local_dl_params(dl_param_id)      
        if error_dl_params == True:
            dl_results = []
            dl_result = {
            'id': 0,
            'dl_result':'-1000',
            'dl_qval':-1000,
            'dl_info':' Loading local dl parameters was not sucessfull',
            }
            dl_results.append(dl_result)
            return dl_results
        logging_text = ' local DL parameters: '+str(reject_thres)+','+str(max_num_results)+','+str(oc_reject_thres)+','+str(oc_num_single_results)
        if db_type == "File":
            write_log_api_file(module_dir,module,logging_text)  
        else:
            write_log_api(module_dir,session,StatusData,status_datas_schema,module,logging_text)  
 
    start_timer_3 = time.time()
    AC_classes = ''
    dl_results = TensorServingGRPC_Client(x_test, module_dir, db_type, session, StatusData, status_datas_schema, DlInitInfo, dl_init_infos_schema, ip_address_tf, port_tf, ai_type_dl, dl_init_ref, user_id, reject_thres, max_num_results)
    dl_results = dl_results[0]
    end_timer_3 = time.time()
    timer_3 = end_timer_3 - start_timer_3
    os.remove('numpy/x_test.npy')
    logging_text = ' Time for prediction and returning the class value: ' + str(timer_3) + ' seconds'
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)  
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
   
    end_timer_1 = time.time()
    timer_1 = end_timer_1 - start_timer_1
    logging_text = ' Total time for one whole prediction: ' + str(timer_1) + ' seconds'
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)  
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)    
   
    return dl_results
 
#Authentication of REST the service
auth = HTTPBasicAuth()
ac_api_rest = Flask(__name__)
ac_api_rest.secret_key = "SecretKeyTextRecogApi"
ac_api_rest.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
CORS(ac_api_rest)    
 
with open('/home/ac/ai_ac_service/api/ac_auth_param_server.json', 'r') as authParams:
    data = authParams.read()
    obj = json.loads(data)
    server_username = str(obj['username'])
    server_password = str(obj['password'])
 
@auth.get_password
def get_password(username):
    if username == server_username:
        return server_password
    return None
 
@auth.error_handler
def unauthorized():
    return make_response(jsonify({'error': 'Unauthorized access'}), 401)
 
@ac_api_rest.route('/ac/api/set_logging', methods=['POST'])
@auth.login_required
def set_logging_status():
   
    module = 'set_logging '
 
    # set additional logging to OFF : 17/9/24 - to be re-considered
    add_log_status = 'OFF'
 
    start_time = time.time()
   
    if not request.json or not 'log_status' in request.json:
        abort(400)
    log_status = request.json['log_status']
    logging_text = ' setting logging status to: '+log_status
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)  
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
   
    if not request.json or not 'log_status_info' in request.json:
        abort(400)
    log_status_info = request.json['log_status_info']
    logging_text = ' setting logging status info to: '+log_status_info
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)  
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
   
    if not request.json or not 'user_id' in request.json:
        abort(400)
    user_id = request.json['user_id']
    logging_text = ' setting user_id to: '+user_id
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)  
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
 
    if db_type == 'File':
        # call the function in basic_functions_cython.pyx file
        Error = set_global_file_data(module_dir, log_status, add_log_status)
        if Error:
            SetStatusOk = False
        else:
            SetStatusOk = True
    else:
        SetStatusOk = set_init_status(session, StatusData, status_datas_schema, '', '', '', '', '', '', log_status, log_status_info, user_id)
    if SetStatusOk:
        status = 'OK'
        status_info = 'setting logging status was successful'
    else:
        status = 'NOK'
        status_info = 'setting logging status was not successful'
   
    end_time = time.time()
    time_taken = end_time-start_time
   
    logging_text = ' logging status: '+ status + ', ' + status_info + ', Time for set logging function: ' + str(time_taken) + ' seconds'
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)  
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
 
    response = {'log status': status}, {'log info': status_info}
    return jsonify(response)
 
 
@ac_api_rest.route('/ac/api/extract_license_number', methods=['POST'])
@auth.login_required
def extract_license_number():
    # TODO: add logging
    module = 'extract_license_number'
 
    try:
        data = request.json
        file_path = data['path'] # type: ignore
        logging_text = 'file path: ' + file_path
        if db_type == 'File':
            write_log_api_file(module_dir, module, logging_text)  
        else:
            write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
    except (BadRequest, TypeError, KeyError) as e:
        logging_text = "Error: " + str(e)
        if db_type == 'File':
            write_log_api_file(module_dir, module, logging_text)  
        else:
            write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
        abort(400)
 
    if not os.path.exists(file_path):
        logging_text = f"Error: {file_path} does not exist"
        if db_type == 'File':
            write_log_api_file(module_dir, module, logging_text)  
        else:
            write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
        return jsonify({'error' : 'Invalid file path'})
 
    if (result := extract_license(file_path)) is None:
        return jsonify({'error' : 'Invalid file type/file does not contain license number'})
 
    return jsonify({ 'license_number' : result })
 
@ac_api_rest.route('/ac/api/get_version_info', methods=['POST'])
@auth.login_required
def get_version_info():
   
    module = 'get_version_info '
   
    version_infos = []
    with open('version_info_ac.json', 'r') as jsonParamsFile:
        data = jsonParamsFile.read()
        obj = json.loads(data)
        service_type = str(obj['service_type'])
        version = str(obj['version'])
        version_num_major = int(obj['version_num_major'])
        version_num_minor = int(obj['version_num_minor'])
        version_num_bugfix = int(obj['version_num_bugfix'])
   
    version_info = {'service_type': service_type, 'version': version, 'version_num_major': version_num_major, 'version_num_minor': version_num_minor, 'version_num_bugfix': version_num_bugfix}
    version_infos.append(version_info)
 
    if db_type == 'File':
        SetVersionOk = True
    else:
        SetVersionOk = set_version_info(session, VersionInfo, version_infos_schema, service_type, version, version_num_major, version_num_minor, version_num_bugfix)
   
    if SetVersionOk:
        logging_text = 'version info: ' + str(version_infos)
        if db_type == 'File':
            write_log_api_file(module_dir, module, logging_text)  
        else:
            write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
    else:
        logging_text = 'set version info in DB was not successful'
        if db_type == 'File':
            write_log_api_file(module_dir, module, logging_text)  
        else:
            write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
   
    with open('version_info_basic.json', 'r') as jsonParamsFile:
        data = jsonParamsFile.read()
        obj = json.loads(data)
        service_type = str(obj['service_type'])
        version = str(obj['version'])
        version_num_major = int(obj['version_num_major'])
        version_num_minor = int(obj['version_num_minor'])
        version_num_bugfix = int(obj['version_num_bugfix'])
   
    version_info = {'service_type': service_type, 'version': version, 'version_num_major': version_num_major, 'version_num_minor': version_num_minor, 'version_num_bugfix': version_num_bugfix}
    version_infos.append(version_info)
 
    if db_type == 'File':
        SetVersionOk = True
    else:
        SetVersionOk = set_version_info(session, VersionInfo, version_infos_schema, service_type, version, version_num_major, version_num_minor, version_num_bugfix)
   
    if SetVersionOk:
        logging_text = 'version info: '+ str(version_infos)
        if db_type == 'File':
            write_log_api_file(module_dir, module, logging_text)  
        else:
            write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
    else:                                                                                                                              
        logging_text = 'set version info in DB was not successful'
        if db_type == 'File':
            write_log_api_file(module_dir, module, logging_text)  
        else:
            write_log_api(module_dir,session, StatusData, status_datas_schema, module, logging_text)
   
    response = {'version_infos': version_infos}
    return jsonify(response)                                                                                  
 
@ac_api_rest.route('/ac/api/set_dl_init_param', methods=['POST'])
@auth.login_required
def set_dl_init_param():
   
    module = 'set_dl_init_param '
   
    start_time = time.time()
   
    if not request.json or not 'description' in request.json:
        abort(400)
    description = request.json['description']
    logging_text = ' description: '+str(description)
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
   
    if not request.json or not 'reject_thres' in request.json:
        abort(400)
    reject_thres = request.json['reject_thres']
    logging_text = ' setting DL reject threshold to: '+str(reject_thres)
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
   
    if not request.json or not 'max_num_results' in request.json:
        abort(400)
    max_num_results = request.json['max_num_results']
    logging_text = ' setting DL result dimension to: '+str(max_num_results)
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
 
    if not request.json or not 'oc_reject_thres' in request.json:
        abort(400)
    oc_reject_thres = request.json['oc_reject_thres']
    logging_text = ' setting DL oc_reject threshold to: '+str(oc_reject_thres)
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
   
    if not request.json or not 'oc_num_single_results' in request.json:
        abort(400)
    oc_num_single_results = request.json['oc_num_single_results']
    logging_text = ' setting DL result dimension to: '+str(oc_num_single_results)
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
 
    Error =  write_dl_init_param(module_dir, db_type, session, StatusData, status_datas_schema, reject_thres, max_num_results, oc_reject_thres, oc_num_single_results)
    if Error == False:
        status = 'OK'
        status_info = 'setting set_dl_init_param was successful'
    else:
        status = 'NOK'
        status_info = 'setting set_dl_init_param was not successful'
    logging_text = ' set_dl_init_param status: '+status+', '+status_info
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)  
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
   
    end_time = time.time()
    time_taken = end_time - start_time
   
    logging_text = ' Time for global parmeter setting operation: ' + str(time_taken) + ' seconds'
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)  
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
   
    response = {'status':status, 'status_info':status_info}
    return jsonify(response)
 
@ac_api_rest.route('/ac/api/set_dl_param', methods=['POST'])
@auth.login_required
def set_dl_param():
 
    module = 'set_dl_param'
 
    start_time = time.time()
 
    if not request.json or not 'description' in request.json:
        abort(400)
    description = request.json['description']
    logging_text = ' description: '+str(description)
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
       
    if not request.json or not 'reject_thres' in request.json:
        abort(400)
    reject_thres = request.json['reject_thres']
    logging_text = ' reject_thres : '+str(reject_thres)
    if db_type == "File":
        write_log_api_file(module_dir,module,logging_text)  
    else:
        write_log_api(module_dir,session,StatusData,status_datas_schema,module,logging_text)
 
    if not request.json or not 'max_num_results' in request.json:
        abort(400)
    max_num_results = request.json['max_num_results']
    logging_text = ' max_num_results : '+str(max_num_results)
    if db_type == "File":
        write_log_api_file(module_dir,module,logging_text)  
    else:
        write_log_api(module_dir,session,StatusData,status_datas_schema,module,logging_text)
 
    if not request.json or not 'oc_reject_thres' in request.json:
        abort(400)
    oc_reject_thres = request.json['oc_reject_thres']
    logging_text = ' oc_reject_thres : '+str(oc_reject_thres)
    if db_type == "File":
        write_log_api_file(module_dir,module,logging_text)  
    else:
        write_log_api(module_dir,session,StatusData,status_datas_schema,module,logging_text)
 
    if not request.json or not 'oc_num_single_results' in request.json:
        abort(400)
    oc_num_single_results = request.json['oc_num_single_results']
    logging_text = ' oc_num_single_results : '+str(oc_num_single_results)
    if db_type == "File":
        write_log_api_file(module_dir,module,logging_text)  
    else:
        write_log_api(module_dir,session,StatusData,status_datas_schema,module,logging_text)
 
    dl_param_id = set_dl_param_code(description, reject_thres, max_num_results, oc_reject_thres, oc_num_single_results)
 
    end_time = time.time()
    time_taken = end_time - start_time
   
    logging_text = ' Time for local parameter setting operation: ' + str(time_taken) + ' seconds'
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)  
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
 
    response = {'dl_param_id':dl_param_id}
   
    return jsonify(response)
 
@ac_api_rest.route('/ac/api/get_dl_param', methods=['POST'])
@auth.login_required
def get_dl_param():
 
    module = 'get_dl_param '
   
    start_time = time.time()
   
    if not request.json or not 'param_list' in request.json:
        abort(400)
    param_list = request.json['param_list']
   
    resp = get_dl_param_code(param_list)
 
    end_time = time.time()
    time_taken = end_time - start_time
   
    logging_text = ' Time for local parameter get operation: ' + str(time_taken) + ' seconds'
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)  
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
 
    return jsonify(resp)
 
@ac_api_rest.route('/ac/api/del_dl_param', methods=['POST'])
@auth.login_required
def del_dl_param():
 
    module = 'del_dl_param '
   
    start_time = time.time()
   
    if not request.json or not 'param_id' in request.json:
        abort(400)
    param_id = request.json['param_id']
    logging_text = ' param_id: '+str(param_id)
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
 
    resp = del_dl_param_code(param_id)
 
    end_time = time.time()
    time_taken = end_time - start_time
   
    logging_text = ' Time for local parameter delete operation: ' + str(time_taken) + ' seconds'
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)  
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
 
    return jsonify(resp)
 
 
@ac_api_rest.route('/ac/api/dl_init', methods=['POST'])
@auth.login_required
def dl_init():
   
    module = 'dl_init'
 
    start_time = time.time()
 
    if not request.json or not 'ts_model_name' in request.json:
        abort(400)
    ts_model_name = request.json['ts_model_name']
    logging_text = ' tensor serve model name: '+ts_model_name
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
   
    if not request.json or not 'ts_model_version' in request.json:
        abort(400)
    ts_model_version = request.json['ts_model_version']
    logging_text = ' tensor serve model version: '+ts_model_version
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
 
    if not request.json or not 'userid' in request.json:
        abort(400)
    user_id = request.json['userid']
    logging_text = ' user id for dl_init: '+user_id
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
   
    if not request.json or not 'dlinit_info' in request.json:
        abort(400)
    dl_init_info = request.json['dlinit_info']
    logging_text = ' init info for Text classifier: '+dl_init_info
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
   
    if not request.json or not 'dl_param_id' in request.json:
        abort(400)
    dl_param_id = request.json['dl_param_id']
    logging_text = ' local params id: '+dl_param_id
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
   
    # create unique init_ref
    init_ref = str(uuid.uuid4())
 
    # load model classes from meta file
    tensor_flow_serving_path = module_dir+'/tfserving/tfservingmodels/'+ts_model_name+'/'
    classes_file_name =  tensor_flow_serving_path + ts_model_name + '_trained_classes.json'
    logging_text = ' Classes model meta data file: '+classes_file_name
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)  
    else:
        write_log_api(module_dir,session,StatusData,status_datas_schema,module,logging_text)
   
    AC_Model_Classes = ''
    if os.path.exists(classes_file_name):
        with open(classes_file_name, 'r') as myfile:
            classes=json.load(myfile)
        try:
            for i in range(0,len(classes)):
                if i < (len(classes)-1):
                    AC_Model_Classes = AC_Model_Classes + classes[i]['class'] + ';'
                else:
                    AC_Model_Classes = AC_Model_Classes + classes[i]['class']
        except:
            AC_Model_Classes = ''
 
    if AC_Model_Classes == '':
        dl_init_status = 'NOK'
        dl_init_status_info = 'reading classes from model meta data file was not successful'
        logging_text = ' DL init status: '+dl_init_status+', '+dl_init_status_info
        if db_type == "File":
            write_log_api_file(module_dir, module, logging_text)  
        else:
            write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
    else:
        dl_init_status = 'OK'
        dl_init_status_info = 'reading classes from model meta data file was successful'
        logging_text = ' DL init status: '+dl_init_status+', '+dl_init_status_info
        if db_type == "File":
            write_log_api_file(module_dir, module, logging_text)  
        else:
            write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
 
    if db_type == "File":
        SetDlInitInfoOK = set_dl_init_info_file(module_dir, init_ref, dl_init_info, ts_model_name, ts_model_version, AC_Model_Classes, user_id, dl_param_id)
    else:
        SetDlInitInfoOK = set_dl_init_info(session, DlInitInfo, init_ref, dl_init_info, ts_model_name, ts_model_version, AC_Model_Classes, user_id)
    if SetDlInitInfoOK:
        dl_init_status = 'OK'
        dl_init_status_info = 'setting DL Init Info in DB was successful'
    else:
        dl_init_status = 'NOK'
        dl_init_status_info = 'setting DL Init Info in DB was not successful'
    logging_text = ' DL Init Info status: '+dl_init_status+', '+dl_init_status_info
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)  
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
 
    if db_type != "File":
        # set init data
        DlInitDataOk =  set_dl_init_data(module_dir, session, StatusData, status_datas_schema, DlInit, ts_model_name, ts_model_version, AC_Model_Classes)
        if DlInitDataOk:
            dl_init_status = 'OK'
            dl_init_status_info = 'setting DL Init Data in DB was successful'
        else:
            dl_init_status = 'NOK'
            dl_init_status_info = 'setting DL Init Data in DB was not successful'
        logging_text = ' DL Init Data status: '+dl_init_status+' '+dl_init_status_info
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
 
    if db_type != "File":
        # set status data
        SetStatusOk = set_init_status(session, StatusData, status_datas_schema, '', '', '', '', dl_init_status, dl_init_status_info, '', '', user_id)
        if SetStatusOk:
            dl_init_status = 'OK'
            dl_init_status_info = 'setting DL Init Status in DB was successful'
        else:
            dl_init_status = 'NOK'
            dl_init_status_info = 'setting DL Init Status in DB was not successful'
        logging_text = ' DL Init status: '+dl_init_status+', '+dl_init_status_info
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
 
    end_time = time.time()
    time_taken = end_time - start_time
   
    logging_text = ' Time for ac_dl_init operation: ' + str(time_taken) + ' seconds'
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)  
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
 
    response = {'init_ref': init_ref, 'init_status': dl_init_status, 'init_info': dl_init_status_info}
    return jsonify(response)
 
@ac_api_rest.route('/ac/api/dl_get_text_result', methods=['POST'])
@auth.login_required
def dl_get_text_result():
   
    module = 'dl_get_text_result'
   
    start_time = time.time()
   
    if not request.json or not 'text_content' in request.json:
        abort(400)
    text_content = request.json['text_content']
   
    if not request.json or not 'ip_address_tf' in request.json:
        abort(400)
    ip_address_tf = request.json['ip_address_tf']
    logging_text = ' IP address TensorFlow Serving: '+ip_address_tf
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
   
    if not request.json or not 'port_tf' in request.json:
        abort(400)
    port_tf = request.json['port_tf']
    logging_text = ' Port TensorFlow Serving: '+port_tf
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
   
    if not request.json or not 'dl_init_ref' in request.json:
        abort(400)
    dl_init_ref = request.json['dl_init_ref']
    logging_text = ' init ref for DL: '+dl_init_ref
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
   
    if not request.json or not 'user_id' in request.json:
        abort(400)
    user_id = request.json['user_id']
    logging_text = ' user id for dl_get_text_result: '+user_id
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
 
    if not request.json or not 'dl_param_id' in request.json:
        abort(400)
    dl_param_id = request.json['dl_param_id']
    logging_text = ' DL param id for dl_get_text_result: '+dl_param_id
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
   
    end_time = time.time()
    time_taken = end_time-start_time
 
    logging_text = ' Time for extracting the data sent from client side: ' + str(time_taken) + ' seconds'
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
 
    response = dl_get_text_result_body(module, text_content, ip_address_tf, port_tf, dl_init_ref, user_id, dl_param_id)
    return jsonify(response)
 
@ac_api_rest.route('/ac/api/dl_get_result', methods=['POST'])
@auth.login_required
def dl_get_result():
   
    module = 'dl_get_result'
   
    start_time = time.time()
   
    if not request.json or not 'binary_content' in request.json:
        abort(400)
    binary_content = request.json['binary_content']
   
    if not request.json or not 'ip_address_tf' in request.json:
        abort(400)
    ip_address_tf = request.json['ip_address_tf']
    logging_text = ' IP address TensorFlow Serving: '+ip_address_tf
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)  
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
   
    if not request.json or not 'port_tf' in request.json:
        abort(400)
    port_tf = request.json['port_tf']
    logging_text = ' Port TensorFlow Serving: '+port_tf
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)  
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
   
    if not request.json or not 'dl_init_ref' in request.json:
        abort(400)
    dl_init_ref = request.json['dl_init_ref']
    logging_text = ' init ref for DL: '+dl_init_ref
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)  
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
   
    if not request.json or not 'suffix' in request.json:
        abort(400)
    suffix = request.json['suffix']
    logging_text = ' file suffix: '+suffix
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)  
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
   
    if not request.json or not 'user_id' in request.json:
        abort(400)
    user_id = request.json['user_id']
    logging_text = ' user id for dl_get_result: '+user_id
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
   
    if not request.json or not 'dl_param_id' in request.json:
        abort(400)
    dl_param_id = request.json['dl_param_id']
    logging_text = ' DL param id for dl_get_text_result: '+dl_param_id
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
 
    end_time = time.time()
    time_taken = end_time-start_time
 
    logging_text = ' Time for receiving data from client side: ' + str(time_taken) + ' seconds'
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)  
    else:
        write_log_api(module_dir ,session, StatusData, status_datas_schema, module, logging_text)
 
    binary_content = base64.b64decode(binary_content)
    #text_content = function_write_file_extract_data(binary_content, suffix)        # write binary data into corresponding file, load from it and perform predictions
    text_content = function_extract_data_from_byte_stream(binary_content, suffix)   # directly convert the binary data into text data and perform predictions
 
    response = dl_get_text_result_body(module, text_content, ip_address_tf, port_tf, dl_init_ref, user_id, dl_param_id)
    return jsonify(response)
 
@ac_api_rest.route('/ac/api/ac_dl_un_init', methods=['POST'])
@auth.login_required
def ac_dl_un_init():
   
    module = 'ac_dl_un_init '
 
    start_time = time.time()
 
    if not request.json or not 'init_ref' in request.json:
        abort(400)
    init_ref = request.json['init_ref']
    logging_text = ' init_ref: '+init_ref
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
 
    if not request.json or not 'user_id' in request.json:
        abort(400)
    user_id = request.json['user_id']
    logging_text = ' user id for ac_dl_un_init: '+user_id
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
   
    init_ref_list = []
    init_ref_list.append(init_ref)
 
    if db_type == "File":
        DelDlInitDataOK, num_files_deleted = delete_dl_info_data_file(module_dir, init_ref_list)
    else:
        DelDlInitDataOK = True  # this must be added to the DB functions if needed
 
    if DelDlInitDataOK:
        status = 'OK'
        status_info  = 'dl init data have been deleted sucessfully'
    else:
        status = 'NOK'
        status_info  = 'no dl init data have been deleted'
 
    logging_text = 'delete dl init data: '+status+', '+status_info
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)  
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
 
    deletion_status = status
    deletion_status_info = status_info
 
    # set StatusData data
    if db_type == "File":
        SetStatusDataOk = True
    else:
        SetStatusDataOk = set_init_status(session, StatusData, status_datas_schema, status, status_info, '', '', '', '', '', '', user_id)
 
    if SetStatusDataOk:
        status = 'OK'
        status_info = 'setting StatusData in DB was successful'
    else:
        status = 'NOK'
        status_info = 'setting StatusData in DB was not successful'
 
    logging_text = 'delete dl init data StatusData: '+status+', '+status_info
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)  
    else:
        write_log_api(module_dir, session ,StatusData, status_datas_schema, module, logging_text)
 
    response = {'init_ref': init_ref, 'init_status':deletion_status, 'init_info':deletion_status_info}
    return jsonify(response)
 

# Function to read init parameters
def read_init_param(set_module_dir):
    # Static initialization values
    ai_type = "DL"
    ai_type_dl = "DL_Class"
    ai_type_knn = "kNN_Class"
    db_type = "File"
    num_threads_gprc_server= "server"
    
    # Return the parameters
    return ai_type, ai_type_dl, ai_type_knn, db_type, num_threads_gprc_server

# set AI module
ai_module = 'ac'
module = 'ac_api_rest'
 
module_dir = set_module_dir(ai_module)
print('Training/reference data path: ', module_dir)
ai_type, ai_type_dl, ai_type_knn, db_type, num_threads_gprc_server = read_init_param(module_dir)
print('db_type: ', db_type) 
#  print or log these values to verify
print(f"AI Type: {ai_type}")
print(f"AI Type (DL): {ai_type_dl}")
print(f"AI Type (KNN): {ai_type_knn}")
print(f"Database Type: {db_type}")

 
if db_type == "File":
   
    session = 0
    VersionInfo = 0
    version_infos_schema = 0
    DlInit = 0
    dl_inits_schema = 0
    StatusData = 0
    status_datas_schema = 0
    DlInitInfo = 0
    dl_init_infos_schema = 0
 
    Error = init_global_file_data(module_dir)
    if Error == True:
        logging_text = ' Error Init Global File Data'
        write_log_api_file(module_dir, module, logging_text)
        print(module+logging_text)
    else:
        logging_text = 'File system will be used instead of database'
        write_log_api_file(module_dir, module, logging_text)
        print(module+logging_text)
        logging_text = 'AI Init Values: '+ai_type+' '+ai_type_dl+' '+ai_type_knn+' '+db_type
        write_log_api_file(module_dir, module, logging_text)
        print(module, logging_text)
else:
    pg_port = '5435'        
    Base, session, init_data = db_session(module_dir, db_type, pg_port)
 
    VersionInfo = init_data[0]
    version_infos_schema = init_data[1]
    DlInit = init_data[8]
    dl_inits_schema = init_data[9]
    StatusData = init_data[10]
    status_datas_schema = init_data[11]
    DlInitInfo = init_data[16]
    dl_init_infos_schema = init_data[17]
 
    logging_text = 'AI Init Values: '+ai_type+' '+ai_type_dl+' '+ai_type_knn+' '+db_type
    write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
 
if __name__ == '__main__':
    #ac_api_rest.run(host="localhost", port=6350, debug=True)                                                          # for insecure communication
    ac_api_rest.run(host="localhost", port=6350, debug=True, ssl_context=('/home/ac/ai_ac_service/api/server_local.crt', '/home/ac/ai_ac_service/api/server_local.key'))     # for secure communication



############################################################################################################################


import os
import io
import sys
import ast
import uuid
import json
import grpc
import time
import base64
import logging
import requests
import docx2txt 
import numpy as np
import pickle as pl
from pracs import map_func
#from pdfcode import *
from docx import Document
from flask_cors import CORS
from datetime import datetime
from datetime import timedelta
from pdfreader import SimplePDFViewer
from flask_httpauth import HTTPBasicAuth
from flask import Flask, jsonify, abort, request, make_response, Response
from werkzeug.exceptions import BadRequest
from filelock import FileLock 
from ac_base import set_dl_param_code, get_dl_param_code, del_dl_param_code
from ac_base import read_local_dl_params, extract_license
from dotenv import load_dotenv 
from pydantic_settings import BaseSettings
from typing import Optional

app = Flask(__name__)

# Load environment variables from the .env files before using them
load_dotenv(dotenv_path= "/home/ac/ai_ac_service/api/init_param.env")  # Load the first .env file
load_dotenv(dotenv_path="/home/ac/ai_ac_service/api/init_dl_param.env")  # Load the second .env file

# Function to print environment variables after loading
def print_environment_variables():
    print("\nLoaded Environment Variables:")
    for key, value in os.environ.items():
        print(f"{key}: {value}")
        
# Print the environment variables after loading them
print_environment_variables()

# Define a Pydantic model for reading general initialization parameters
class InitParams(BaseSettings):
    ai_type: str
    ai_type_dl: str
    ai_type_knn: str
    db_type: str

# Define a Pydantic model for reading deep learning related initialization parameters
class DLInitParams(BaseSettings):
    description: str
    reject_thres: float
    max_num_results: int
    oc_reject_thres: float
    oc_num_single_results: int

# Load the environment variables using Pydantic models
params = InitParams()
dl_params = DLInitParams()

# function/method for checking if this code is running within docker
def is_docker():
    path = '/proc/self/cgroup'
    return (
        os.path.exists('/.dockerenv') or
        os.path.isfile(path) and any('docker' in line for line in open(path))
    )

# Rest of your code continues...


if is_docker():
    sys.path.insert(0, '../basicdata/')
else:
    # get current directory of AI service
    curr_dir = os.getcwd()
    print('Current AI service directory: ', curr_dir)
    # get basic directory
    basic_dir = curr_dir+'/basicdata/'
    print('Basic directory: ', basic_dir)
    sys.path.insert(0, basic_dir)

def token_func(ts_model_name):
    if is_docker():
        tokenize_path = '../models/'+ts_model_name+'/tokenizer.json'
    else:
        tokenize_path = './tfserving/tfservingmodels/'+ts_model_name+'/tokenizer.json'
    with open(tokenize_path) as f:
        data = json.load(f)
    return data 

 
 
# Flask route for testing the loaded parameters
@app.route('/test_params', methods=['GET'])
def test_params():
    response = {
        "description": params.description,
        "ai_type": params.ai_type,
        "ai_type_dl": params.ai_type_dl,
        "ai_type_knn": params.ai_type_knn,
        "db_type": params.db_type,
        "reject_thres": dl_params.reject_thres,
        "max_num_results": dl_params.max_num_results,
        "oc_reject_thres": dl_params.oc_reject_thres,
        "oc_num_single_results": dl_params.oc_num_single_results
    }
    return jsonify(response)

from basic_functions_cython import set_module_dir
from basic_functions_cython import write_log_api
from basic_functions_cython import write_log_api_file
from basic_functions_cython import init_global_file_data
# from basic_functions_cython import read_init_param  
from basic_functions_cython import write_dl_init_param
from basic_functions_cython import set_dl_init_info_file
from basic_functions_cython import get_dl_init_info_file
# from basic_functions_cython import read_dl_init_param
from basic_functions_cython import delete_dl_info_data_file
from basic_functions_cython import TensorServingGRPC_Client
from basic_functions_cython import set_global_file_data
from basic_db_functions_cython import db_session
from basic_db_functions_cython import get_init_status
from basic_db_functions_cython import set_init_status
from basic_db_functions_cython import set_dl_init_data
from basic_db_functions_cython import set_version_info
from basic_db_functions_cython import set_dl_init_info
from basic_db_functions_cython import get_dl_init_data
 
def function_write_file_extract_data(binary_content, suffix):
    unique_time_stamp = str(int(time.time()))
    filename = unique_time_stamp+suffix
    data_file = open(filename, "wb")
    data_file.write(binary_content)
    data_file.close()
    if os.path.isfile(filename):
        data_file = open((filename), "r")
        file = data_file.name
        if file.endswith('.docx'):
            extracted_string_data = docx2txt.process(data_file.name)
        elif file.endswith('.pdf'):
            logging_text = 'Please do not use PDF file'
            if db_type == "File":
                write_log_api_file(module_dir, module, logging_text)  
            else:
                write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
            #extracted_string_data = Convert.convert_pdf_to_txt(data_file.name)
            extracted_string_data = 'No Data'
        elif file.endswith('.txt'):
            extracted_string_data = data_file.read()
        else:
            extracted_string_data = data_file.read()
    else:
        print('file not found')
    os.remove(filename)
    new_string = extracted_string_data
    chars_to_remove = [ "@", "1", "2", "3", "4", "5", "6", "7", "8", "9", "0", ".", ",", ";", ":", "-", "\r", "\n", ">", "<"]
    for element in new_string:
        if element in chars_to_remove:
            new_string = new_string.replace(element, ' ')
    special_chars_to_remove = ["&"]
    for element in new_string:
        if element in special_chars_to_remove:
            index_special_char = int(new_string.index('&'))
            if index_special_char==0:
                new_string = ' ' + new_string[3 : : ]
            elif index_special_char==len(new_string)-2:
                new_string = new_string.replace(element, ' ')
            elif index_special_char==len(new_string)-1:
                new_string = new_string.replace(element, ' ')
            else:
                start = index_special_char
                stop = index_special_char+2
                new_string = new_string[0: start:] + ' ' + new_string[stop + 1::]
    extracted_string_data = new_string
    return extracted_string_data
 
def function_extract_data_from_byte_stream(binary_content, suffix):
    extracted_string_data=''
    if suffix=='pdf':
        viewer = SimplePDFViewer(binary_content)
        viewer.render()
        string_data = viewer.canvas.strings
        string_data = ''.join(string_data)
        extracted_string_data = string_data
    elif suffix=='docx':
        string_data = ''
        doc = Document(io.BytesIO(binary_content))
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        string_data = ''.join(full_text)
        extracted_string_data = string_data
    else:
        string_data = binary_content.decode('utf-8')
        extracted_string_data = string_data
    new_string = extracted_string_data
    chars_to_remove = [ "@", "1", "2", "3", "4", "5", "6", "7", "8", "9", "0", ".", ",", ";", ":", "-", "\r", "\n", ">", "<"]
    for element in new_string:
        if element in chars_to_remove:
            new_string = new_string.replace(element, ' ')
    special_chars_to_remove = ["&"]
    for element in new_string:
        if element in special_chars_to_remove:
            index_special_char = int(new_string.index('&'))
            if index_special_char==0:
                new_string = ' ' + new_string[3 : : ]
            elif index_special_char==len(new_string)-2:
                new_string = new_string.replace(element, ' ')
            elif index_special_char==len(new_string)-1:
                new_string = new_string.replace(element, ' ')
            else:
                start = index_special_char
                stop = index_special_char+2
                new_string = new_string[0: start:] + ' ' + new_string[stop + 1::]
    extracted_string_data = new_string
    return extracted_string_data
 
def dl_get_text_result_body(module, text_content, ip_address_tf, port_tf, dl_init_ref, user_id, dl_param_id=''):
   
    module = module + '/dl_get_text_result_body '
   
    start_timer_1 = time.time()
 
    if db_type == 'File':
        status_current = 'OK'
    else:
        status_data = get_init_status(session, StatusData, status_datas_schema, user_id)
        if len(status_data) > 0:
            status_current = status_data['status']
            status_info_current = status_data['status_info']
            dl_init_status_current = status_data['dl_init_status']
            dl_init_status_info_current = status_data['dl_init_status_info']
        else:
            dl_results = []
            dl_result = {'id': 0, 'dl_result':'-1000', 'dl_qval':'-1000', 'dl_info':'DL init status not available'}
            dl_results.append(dl_result)
            return jsonify(dl_results)
 
        if dl_init_status_current == 'NOK':
            dl_results = []
            dl_result = {'id': 0, 'dl_result':'-1000', 'dl_qval':-1000, 'dl_info':'DL classifier not initialized'}
            dl_results.append(dl_result)
            return jsonify(dl_results)
 
    if status_current == 'OK':
        if db_type == "File":
            GetDlInitDataInfoOK, dl_init_data_info_list = get_dl_init_info_file(module_dir, dl_init_ref)
            if GetDlInitDataInfoOK:
                ts_model_name = dl_init_data_info_list[0].get('dlinit_tsmodelname')
                ts_model_version = dl_init_data_info_list[0].get('dlinit_tsmodelversion')
                ts_model_classes = dl_init_data_info_list[0].get('dlinit_tsmodelclasses')
                logging_text = ' TF Serving Model, Version & Classes successfully loaded from data/dl_init_info_xxxxxxxx.json file'
                write_log_api_file(module_dir, module, logging_text)
            else:
                logging_text = ' TF Serving Model, Version & Classes could not be successfully loaded from data/dl_init_info_xxxxxxxx.json file'
                write_log_api_file(module_dir, module, logging_text)
                ts_model_name = ''
                ts_model_version = ''
                ts_model_classes = ''
        else:
            ts_model_name, ts_model_version, ts_model_classes = get_dl_init_data(session, DlInit, dl_inits_schema)
    else:
        dl_results = []
        dl_result = {'id': 0, 'dl_result':'-1000', 'dl_qval':'-1000', 'dl_info':'DL not initialized'}
        dl_results.append(dl_result)
        return jsonify(dl_results)
   
    start_timer_2 = time.time()
    headers = {'Content-Type': 'application/json'}
    tokens= token_func(ts_model_name)
    string_data = {"dirAddress":text_content, "tokenizer": tokens}
 
    #preparing URL for connecting to ai_ac_datapreprocess_api_rest container
    if is_docker():
        ## this works for DOCKER server
        ip_dp='ai_datapreprocess_api_rest'
    else:
        ## this work for MANUAL server
        ip_dp='127.0.0.1'
 
    dp_url='https://'+ip_dp+':6355/ac/api/data_preprocessing'
   
    # following line is for insecure communication
    #classify= requests.post(dp_url.replace('ps:','p:'), headers=headers, json=string_data, auth=(server_username, server_password))
   
    # following line is for secure communication
    classify= requests.post(dp_url, headers=headers, json=string_data, auth=(server_username, server_password), cert=('server_local.crt', 'server_local.key'), verify=False)
     
    end_timer_2 = time.time()
    timer_2 = end_timer_2 - start_timer_2
    logging_text = ' Time for data preprocessing from another service: ' + str(timer_2) + ' seconds'
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)  
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
   
    x_test = np.load('numpy/x_test.npy')
   
    # read_init => need  to remove and read from dotenv file
    #read global DL init data
    # reject_thres, max_num_results, oc_reject_thres, oc_num_single_results = read_dl_init_param(module_dir, db_type, session, StatusData, status_datas_schema)
    logging_text = ' DL reject threshold: '+str(reject_thres)
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)  
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
    logging_text = ' DL maximum number of results per decision: '+str(max_num_results)
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)  
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
 
    # read local DL parameters
    if dl_param_id != '0':
        reject_thres,max_num_results,oc_reject_thres,oc_num_single_results,error_dl_params = read_local_dl_params(dl_param_id)      
        if error_dl_params == True:
            dl_results = []
            dl_result = {
            'id': 0,
            'dl_result':'-1000',
            'dl_qval':-1000,
            'dl_info':' Loading local dl parameters was not sucessfull',
            }
            dl_results.append(dl_result)
            return dl_results
        logging_text = ' local DL parameters: '+str(reject_thres)+','+str(max_num_results)+','+str(oc_reject_thres)+','+str(oc_num_single_results)
        if db_type == "File":
            write_log_api_file(module_dir,module,logging_text)  
        else:
            write_log_api(module_dir,session,StatusData,status_datas_schema,module,logging_text)  
 
    start_timer_3 = time.time()
    AC_classes = ''
    dl_results = TensorServingGRPC_Client(x_test, module_dir, db_type, session, StatusData, status_datas_schema, DlInitInfo, dl_init_infos_schema, ip_address_tf, port_tf, ai_type_dl, dl_init_ref, user_id, reject_thres, max_num_results)
    dl_results = dl_results[0]
    end_timer_3 = time.time()
    timer_3 = end_timer_3 - start_timer_3
    os.remove('numpy/x_test.npy')
    logging_text = ' Time for prediction and returning the class value: ' + str(timer_3) + ' seconds'
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)  
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
   
    end_timer_1 = time.time()
    timer_1 = end_timer_1 - start_timer_1
    logging_text = ' Total time for one whole prediction: ' + str(timer_1) + ' seconds'
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)  
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)    
   
    return dl_results
 
#Authentication of REST the service
auth = HTTPBasicAuth()
ac_api_rest = Flask(__name__)
ac_api_rest.secret_key = "SecretKeyTextRecogApi"
ac_api_rest.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
CORS(ac_api_rest)    
 
with open('/home/ac/ai_ac_service/api/ac_auth_param_server.json', 'r') as authParams:
    data = authParams.read()
    obj = json.loads(data)
    server_username = str(obj['username'])
    server_password = str(obj['password'])
 
@auth.get_password
def get_password(username):
    if username == server_username:
        return server_password
    return None
 
@auth.error_handler
def unauthorized():
    return make_response(jsonify({'error': 'Unauthorized access'}), 401)
 
@ac_api_rest.route('/ac/api/set_logging', methods=['POST'])
@auth.login_required
def set_logging_status():
   
    module = 'set_logging '
 
    # set additional logging to OFF : 17/9/24 - to be re-considered
    add_log_status = 'OFF'
 
    start_time = time.time()
   
    if not request.json or not 'log_status' in request.json:
        abort(400)
    log_status = request.json['log_status']
    logging_text = ' setting logging status to: '+log_status
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)  
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
   
    if not request.json or not 'log_status_info' in request.json:
        abort(400)
    log_status_info = request.json['log_status_info']
    logging_text = ' setting logging status info to: '+log_status_info
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)  
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
   
    if not request.json or not 'user_id' in request.json:
        abort(400)
    user_id = request.json['user_id']
    logging_text = ' setting user_id to: '+user_id
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)  
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
 
    if db_type == 'File':
        # call the function in basic_functions_cython.pyx file
        Error = set_global_file_data(module_dir, log_status, add_log_status)
        if Error:
            SetStatusOk = False
        else:
            SetStatusOk = True
    else:
        SetStatusOk = set_init_status(session, StatusData, status_datas_schema, '', '', '', '', '', '', log_status, log_status_info, user_id)
    if SetStatusOk:
        status = 'OK'
        status_info = 'setting logging status was successful'
    else:
        status = 'NOK'
        status_info = 'setting logging status was not successful'
   
    end_time = time.time()
    time_taken = end_time-start_time
   
    logging_text = ' logging status: '+ status + ', ' + status_info + ', Time for set logging function: ' + str(time_taken) + ' seconds'
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)  
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
 
    response = {'log status': status}, {'log info': status_info}
    return jsonify(response)
 
 
@ac_api_rest.route('/ac/api/extract_license_number', methods=['POST'])
@auth.login_required
def extract_license_number():
    # TODO: add logging
    module = 'extract_license_number'
 
    try:
        data = request.json
        file_path = data['path'] # type: ignore
        logging_text = 'file path: ' + file_path
        if db_type == 'File':
            write_log_api_file(module_dir, module, logging_text)  
        else:
            write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
    except (BadRequest, TypeError, KeyError) as e:
        logging_text = "Error: " + str(e)
        if db_type == 'File':
            write_log_api_file(module_dir, module, logging_text)  
        else:
            write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
        abort(400)
 
    if not os.path.exists(file_path):
        logging_text = f"Error: {file_path} does not exist"
        if db_type == 'File':
            write_log_api_file(module_dir, module, logging_text)  
        else:
            write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
        return jsonify({'error' : 'Invalid file path'})
 
    if (result := extract_license(file_path)) is None:
        return jsonify({'error' : 'Invalid file type/file does not contain license number'})
 
    return jsonify({ 'license_number' : result })
 
@ac_api_rest.route('/ac/api/get_version_info', methods=['POST'])
@auth.login_required
def get_version_info():
   
    module = 'get_version_info '
   
    version_infos = []
    with open('version_info_ac.json', 'r') as jsonParamsFile:
        data = jsonParamsFile.read()
        obj = json.loads(data)
        service_type = str(obj['service_type'])
        version = str(obj['version'])
        version_num_major = int(obj['version_num_major'])
        version_num_minor = int(obj['version_num_minor'])
        version_num_bugfix = int(obj['version_num_bugfix'])
   
    version_info = {'service_type': service_type, 'version': version, 'version_num_major': version_num_major, 'version_num_minor': version_num_minor, 'version_num_bugfix': version_num_bugfix}
    version_infos.append(version_info)
 
    if db_type == 'File':
        SetVersionOk = True
    else:
        SetVersionOk = set_version_info(session, VersionInfo, version_infos_schema, service_type, version, version_num_major, version_num_minor, version_num_bugfix)
   
    if SetVersionOk:
        logging_text = 'version info: ' + str(version_infos)
        if db_type == 'File':
            write_log_api_file(module_dir, module, logging_text)  
        else:
            write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
    else:
        logging_text = 'set version info in DB was not successful'
        if db_type == 'File':
            write_log_api_file(module_dir, module, logging_text)  
        else:
            write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
   
    with open('version_info_basic.json', 'r') as jsonParamsFile:
        data = jsonParamsFile.read()
        obj = json.loads(data)
        service_type = str(obj['service_type'])
        version = str(obj['version'])
        version_num_major = int(obj['version_num_major'])
        version_num_minor = int(obj['version_num_minor'])
        version_num_bugfix = int(obj['version_num_bugfix'])
   
    version_info = {'service_type': service_type, 'version': version, 'version_num_major': version_num_major, 'version_num_minor': version_num_minor, 'version_num_bugfix': version_num_bugfix}
    version_infos.append(version_info)
 
    if db_type == 'File':
        SetVersionOk = True
    else:
        SetVersionOk = set_version_info(session, VersionInfo, version_infos_schema, service_type, version, version_num_major, version_num_minor, version_num_bugfix)
   
    if SetVersionOk:
        logging_text = 'version info: '+ str(version_infos)
        if db_type == 'File':
            write_log_api_file(module_dir, module, logging_text)  
        else:
            write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
    else:                                                                                                                              
        logging_text = 'set version info in DB was not successful'
        if db_type == 'File':
            write_log_api_file(module_dir, module, logging_text)  
        else:
            write_log_api(module_dir,session, StatusData, status_datas_schema, module, logging_text)
   
    response = {'version_infos': version_infos}
    return jsonify(response)                                                                                  
 
@ac_api_rest.route('/ac/api/set_dl_init_param', methods=['POST'])
@auth.login_required
def set_dl_init_param():
   
    module = 'set_dl_init_param '
   
    start_time = time.time()
   
    if not request.json or not 'description' in request.json:
        abort(400)
    description = request.json['description']
    logging_text = ' description: '+str(description)
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
   
    if not request.json or not 'reject_thres' in request.json:
        abort(400)
    reject_thres = request.json['reject_thres']
    logging_text = ' setting DL reject threshold to: '+str(reject_thres)
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
   
    if not request.json or not 'max_num_results' in request.json:
        abort(400)
    max_num_results = request.json['max_num_results']
    logging_text = ' setting DL result dimension to: '+str(max_num_results)
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
 
    if not request.json or not 'oc_reject_thres' in request.json:
        abort(400)
    oc_reject_thres = request.json['oc_reject_thres']
    logging_text = ' setting DL oc_reject threshold to: '+str(oc_reject_thres)
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
   
    if not request.json or not 'oc_num_single_results' in request.json:
        abort(400)
    oc_num_single_results = request.json['oc_num_single_results']
    logging_text = ' setting DL result dimension to: '+str(oc_num_single_results)
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
 
    Error =  write_dl_init_param(module_dir, db_type, session, StatusData, status_datas_schema, reject_thres, max_num_results, oc_reject_thres, oc_num_single_results)
    if Error == False:
        status = 'OK'
        status_info = 'setting set_dl_init_param was successful'
    else:
        status = 'NOK'
        status_info = 'setting set_dl_init_param was not successful'
    logging_text = ' set_dl_init_param status: '+status+', '+status_info
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)  
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
   
    end_time = time.time()
    time_taken = end_time - start_time
   
    logging_text = ' Time for global parmeter setting operation: ' + str(time_taken) + ' seconds'
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)  
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
   
    response = {'status':status, 'status_info':status_info}
    return jsonify(response)
 
@ac_api_rest.route('/ac/api/set_dl_param', methods=['POST'])
@auth.login_required
def set_dl_param():
 
    module = 'set_dl_param'
 
    start_time = time.time()
 
    if not request.json or not 'description' in request.json:
        abort(400)
    description = request.json['description']
    logging_text = ' description: '+str(description)
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
       
    if not request.json or not 'reject_thres' in request.json:
        abort(400)
    reject_thres = request.json['reject_thres']
    logging_text = ' reject_thres : '+str(reject_thres)
    if db_type == "File":
        write_log_api_file(module_dir,module,logging_text)  
    else:
        write_log_api(module_dir,session,StatusData,status_datas_schema,module,logging_text)
 
    if not request.json or not 'max_num_results' in request.json:
        abort(400)
    max_num_results = request.json['max_num_results']
    logging_text = ' max_num_results : '+str(max_num_results)
    if db_type == "File":
        write_log_api_file(module_dir,module,logging_text)  
    else:
        write_log_api(module_dir,session,StatusData,status_datas_schema,module,logging_text)
 
    if not request.json or not 'oc_reject_thres' in request.json:
        abort(400)
    oc_reject_thres = request.json['oc_reject_thres']
    logging_text = ' oc_reject_thres : '+str(oc_reject_thres)
    if db_type == "File":
        write_log_api_file(module_dir,module,logging_text)  
    else:
        write_log_api(module_dir,session,StatusData,status_datas_schema,module,logging_text)
 
    if not request.json or not 'oc_num_single_results' in request.json:
        abort(400)
    oc_num_single_results = request.json['oc_num_single_results']
    logging_text = ' oc_num_single_results : '+str(oc_num_single_results)
    if db_type == "File":
        write_log_api_file(module_dir,module,logging_text)  
    else:
        write_log_api(module_dir,session,StatusData,status_datas_schema,module,logging_text)
 
    dl_param_id = set_dl_param_code(description, reject_thres, max_num_results, oc_reject_thres, oc_num_single_results)
 
    end_time = time.time()
    time_taken = end_time - start_time
   
    logging_text = ' Time for local parameter setting operation: ' + str(time_taken) + ' seconds'
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)  
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
 
    response = {'dl_param_id':dl_param_id}
   
    return jsonify(response)
 
@ac_api_rest.route('/ac/api/get_dl_param', methods=['POST'])
@auth.login_required
def get_dl_param():
 
    module = 'get_dl_param '
   
    start_time = time.time()
   
    if not request.json or not 'param_list' in request.json:
        abort(400)
    param_list = request.json['param_list']
   
    resp = get_dl_param_code(param_list)
 
    end_time = time.time()
    time_taken = end_time - start_time
   
    logging_text = ' Time for local parameter get operation: ' + str(time_taken) + ' seconds'
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)  
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
 
    return jsonify(resp)
 
@ac_api_rest.route('/ac/api/del_dl_param', methods=['POST'])
@auth.login_required
def del_dl_param():
 
    module = 'del_dl_param '
   
    start_time = time.time()
   
    if not request.json or not 'param_id' in request.json:
        abort(400)
    param_id = request.json['param_id']
    logging_text = ' param_id: '+str(param_id)
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
 
    resp = del_dl_param_code(param_id)
 
    end_time = time.time()
    time_taken = end_time - start_time
   
    logging_text = ' Time for local parameter delete operation: ' + str(time_taken) + ' seconds'
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)  
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
 
    return jsonify(resp)
 
 
@ac_api_rest.route('/ac/api/dl_init', methods=['POST'])
@auth.login_required
def dl_init():
   
    module = 'dl_init'
 
    start_time = time.time()
 
    if not request.json or not 'ts_model_name' in request.json:
        abort(400)
    ts_model_name = request.json['ts_model_name']
    logging_text = ' tensor serve model name: '+ts_model_name
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
   
    if not request.json or not 'ts_model_version' in request.json:
        abort(400)
    ts_model_version = request.json['ts_model_version']
    logging_text = ' tensor serve model version: '+ts_model_version
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
 
    if not request.json or not 'userid' in request.json:
        abort(400)
    user_id = request.json['userid']
    logging_text = ' user id for dl_init: '+user_id
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
   
    if not request.json or not 'dlinit_info' in request.json:
        abort(400)
    dl_init_info = request.json['dlinit_info']
    logging_text = ' init info for Text classifier: '+dl_init_info
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
   
    if not request.json or not 'dl_param_id' in request.json:
        abort(400)
    dl_param_id = request.json['dl_param_id']
    logging_text = ' local params id: '+dl_param_id
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
   
    # create unique init_ref
    init_ref = str(uuid.uuid4())
 
    # load model classes from meta file
    tensor_flow_serving_path = module_dir+'/tfserving/tfservingmodels/'+ts_model_name+'/'
    classes_file_name =  tensor_flow_serving_path + ts_model_name + '_trained_classes.json'
    logging_text = ' Classes model meta data file: '+classes_file_name
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)  
    else:
        write_log_api(module_dir,session,StatusData,status_datas_schema,module,logging_text)
   
    AC_Model_Classes = ''
    if os.path.exists(classes_file_name):
        with open(classes_file_name, 'r') as myfile:
            classes=json.load(myfile)
        try:
            for i in range(0,len(classes)):
                if i < (len(classes)-1):
                    AC_Model_Classes = AC_Model_Classes + classes[i]['class'] + ';'
                else:
                    AC_Model_Classes = AC_Model_Classes + classes[i]['class']
        except:
            AC_Model_Classes = ''
 
    if AC_Model_Classes == '':
        dl_init_status = 'NOK'
        dl_init_status_info = 'reading classes from model meta data file was not successful'
        logging_text = ' DL init status: '+dl_init_status+', '+dl_init_status_info
        if db_type == "File":
            write_log_api_file(module_dir, module, logging_text)  
        else:
            write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
    else:
        dl_init_status = 'OK'
        dl_init_status_info = 'reading classes from model meta data file was successful'
        logging_text = ' DL init status: '+dl_init_status+', '+dl_init_status_info
        if db_type == "File":
            write_log_api_file(module_dir, module, logging_text)  
        else:
            write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
 
    if db_type == "File":
        SetDlInitInfoOK = set_dl_init_info_file(module_dir, init_ref, dl_init_info, ts_model_name, ts_model_version, AC_Model_Classes, user_id, dl_param_id)
    else:
        SetDlInitInfoOK = set_dl_init_info(session, DlInitInfo, init_ref, dl_init_info, ts_model_name, ts_model_version, AC_Model_Classes, user_id)
    if SetDlInitInfoOK:
        dl_init_status = 'OK'
        dl_init_status_info = 'setting DL Init Info in DB was successful'
    else:
        dl_init_status = 'NOK'
        dl_init_status_info = 'setting DL Init Info in DB was not successful'
    logging_text = ' DL Init Info status: '+dl_init_status+', '+dl_init_status_info
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)  
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
 
    if db_type != "File":
        # set init data
        DlInitDataOk =  set_dl_init_data(module_dir, session, StatusData, status_datas_schema, DlInit, ts_model_name, ts_model_version, AC_Model_Classes)
        if DlInitDataOk:
            dl_init_status = 'OK'
            dl_init_status_info = 'setting DL Init Data in DB was successful'
        else:
            dl_init_status = 'NOK'
            dl_init_status_info = 'setting DL Init Data in DB was not successful'
        logging_text = ' DL Init Data status: '+dl_init_status+' '+dl_init_status_info
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
 
    if db_type != "File":
        # set status data
        SetStatusOk = set_init_status(session, StatusData, status_datas_schema, '', '', '', '', dl_init_status, dl_init_status_info, '', '', user_id)
        if SetStatusOk:
            dl_init_status = 'OK'
            dl_init_status_info = 'setting DL Init Status in DB was successful'
        else:
            dl_init_status = 'NOK'
            dl_init_status_info = 'setting DL Init Status in DB was not successful'
        logging_text = ' DL Init status: '+dl_init_status+', '+dl_init_status_info
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
 
    end_time = time.time()
    time_taken = end_time - start_time
   
    logging_text = ' Time for ac_dl_init operation: ' + str(time_taken) + ' seconds'
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)  
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
 
    response = {'init_ref': init_ref, 'init_status': dl_init_status, 'init_info': dl_init_status_info}
    return jsonify(response)
 
@ac_api_rest.route('/ac/api/dl_get_text_result', methods=['POST'])
@auth.login_required
def dl_get_text_result():
   
    module = 'dl_get_text_result'
   
    start_time = time.time()
   
    if not request.json or not 'text_content' in request.json:
        abort(400)
    text_content = request.json['text_content']
   
    if not request.json or not 'ip_address_tf' in request.json:
        abort(400)
    ip_address_tf = request.json['ip_address_tf']
    logging_text = ' IP address TensorFlow Serving: '+ip_address_tf
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
   
    if not request.json or not 'port_tf' in request.json:
        abort(400)
    port_tf = request.json['port_tf']
    logging_text = ' Port TensorFlow Serving: '+port_tf
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
   
    if not request.json or not 'dl_init_ref' in request.json:
        abort(400)
    dl_init_ref = request.json['dl_init_ref']
    logging_text = ' init ref for DL: '+dl_init_ref
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
   
    if not request.json or not 'user_id' in request.json:
        abort(400)
    user_id = request.json['user_id']
    logging_text = ' user id for dl_get_text_result: '+user_id
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
 
    if not request.json or not 'dl_param_id' in request.json:
        abort(400)
    dl_param_id = request.json['dl_param_id']
    logging_text = ' DL param id for dl_get_text_result: '+dl_param_id
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
   
    end_time = time.time()
    time_taken = end_time-start_time
 
    logging_text = ' Time for extracting the data sent from client side: ' + str(time_taken) + ' seconds'
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
 
    response = dl_get_text_result_body(module, text_content, ip_address_tf, port_tf, dl_init_ref, user_id, dl_param_id)
    return jsonify(response)
 
@ac_api_rest.route('/ac/api/dl_get_result', methods=['POST'])
@auth.login_required
def dl_get_result():
   
    module = 'dl_get_result'
   
    start_time = time.time()
   
    if not request.json or not 'binary_content' in request.json:
        abort(400)
    binary_content = request.json['binary_content']
   
    if not request.json or not 'ip_address_tf' in request.json:
        abort(400)
    ip_address_tf = request.json['ip_address_tf']
    logging_text = ' IP address TensorFlow Serving: '+ip_address_tf
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)  
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
   
    if not request.json or not 'port_tf' in request.json:
        abort(400)
    port_tf = request.json['port_tf']
    logging_text = ' Port TensorFlow Serving: '+port_tf
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)  
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
   
    if not request.json or not 'dl_init_ref' in request.json:
        abort(400)
    dl_init_ref = request.json['dl_init_ref']
    logging_text = ' init ref for DL: '+dl_init_ref
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)  
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
   
    if not request.json or not 'suffix' in request.json:
        abort(400)
    suffix = request.json['suffix']
    logging_text = ' file suffix: '+suffix
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)  
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
   
    if not request.json or not 'user_id' in request.json:
        abort(400)
    user_id = request.json['user_id']
    logging_text = ' user id for dl_get_result: '+user_id
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
   
    if not request.json or not 'dl_param_id' in request.json:
        abort(400)
    dl_param_id = request.json['dl_param_id']
    logging_text = ' DL param id for dl_get_text_result: '+dl_param_id
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
 
    end_time = time.time()
    time_taken = end_time-start_time
 
    logging_text = ' Time for receiving data from client side: ' + str(time_taken) + ' seconds'
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)  
    else:
        write_log_api(module_dir ,session, StatusData, status_datas_schema, module, logging_text)
 
    binary_content = base64.b64decode(binary_content)
    #text_content = function_write_file_extract_data(binary_content, suffix)        # write binary data into corresponding file, load from it and perform predictions
    text_content = function_extract_data_from_byte_stream(binary_content, suffix)   # directly convert the binary data into text data and perform predictions
 
    response = dl_get_text_result_body(module, text_content, ip_address_tf, port_tf, dl_init_ref, user_id, dl_param_id)
    return jsonify(response)
 
@ac_api_rest.route('/ac/api/ac_dl_un_init', methods=['POST'])
@auth.login_required
def ac_dl_un_init():
   
    module = 'ac_dl_un_init '
 
    start_time = time.time()
 
    if not request.json or not 'init_ref' in request.json:
        abort(400)
    init_ref = request.json['init_ref']
    logging_text = ' init_ref: '+init_ref
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
 
    if not request.json or not 'user_id' in request.json:
        abort(400)
    user_id = request.json['user_id']
    logging_text = ' user id for ac_dl_un_init: '+user_id
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
   
    init_ref_list = []
    init_ref_list.append(init_ref)
 
    if db_type == "File":
        DelDlInitDataOK, num_files_deleted = delete_dl_info_data_file(module_dir, init_ref_list)
    else:
        DelDlInitDataOK = True  # this must be added to the DB functions if needed
 
    if DelDlInitDataOK:
        status = 'OK'
        status_info  = 'dl init data have been deleted sucessfully'
    else:
        status = 'NOK'
        status_info  = 'no dl init data have been deleted'
 
    logging_text = 'delete dl init data: '+status+', '+status_info
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)  
    else:
        write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
 
    deletion_status = status
    deletion_status_info = status_info
 
    # set StatusData data
    if db_type == "File":
        SetStatusDataOk = True
    else:
        SetStatusDataOk = set_init_status(session, StatusData, status_datas_schema, status, status_info, '', '', '', '', '', '', user_id)
 
    if SetStatusDataOk:
        status = 'OK'
        status_info = 'setting StatusData in DB was successful'
    else:
        status = 'NOK'
        status_info = 'setting StatusData in DB was not successful'
 
    logging_text = 'delete dl init data StatusData: '+status+', '+status_info
    if db_type == "File":
        write_log_api_file(module_dir, module, logging_text)  
    else:
        write_log_api(module_dir, session ,StatusData, status_datas_schema, module, logging_text)
 
    response = {'init_ref': init_ref, 'init_status':deletion_status, 'init_info':deletion_status_info}
    return jsonify(response)
 

# Commenting out the read function
# def read_init_param(set_module_dir):
#     # Static initialization values
#     ai_type = "DL"
#     ai_type_dl = "DL_Class"
#     ai_type_knn = "kNN_Class"
#     db_type = "File"
#     num_threads_gprc_server= "server"
#     
#     # Return the parameters
#     return ai_type, ai_type_dl, ai_type_knn, db_type, num_threads_gprc_server

# Integrating dotenv for environment variable management
from dotenv import load_dotenv

# Load environment variables from a .env file
load_dotenv()

# Example of accessing environment variables
#DB_SERVER = os.getenv("DB_SERVER", "default_server")
#DATABASE = os.getenv("DATABASE", "default_database")
#DB_UID = os.getenv("DB_UID", "default_uid")
#DB_PASSWORD = os.getenv("DB_PASSWORD", "default_password")

# set AI module
ai_module = 'ac'
module = 'ac_api_rest'
 
module_dir = set_module_dir(ai_module)
print('Training/reference data path: ', module_dir)

# Replace read_init_param with environment variable loading
ai_type = os.getenv("AI_TYPE", "DL")
ai_type_dl = os.getenv("AI_TYPE_DL", "DL_Class")
ai_type_knn = os.getenv("AI_TYPE_KNN", "kNN_Class")
db_type = os.getenv("DB_TYPE", "File")
num_threads_gprc_server = os.getenv("NUM_THREADS_GPRC_SERVER", "server")

print('db_type: ', db_type) 
#  print or log these values to verify
print(f"AI Type: {ai_type}")
print(f"AI Type (DL): {ai_type_dl}")
print(f"AI Type (KNN): {ai_type_knn}")
print(f"Database Type: {db_type}")
if db_type == "File":
   
    session = 0
    VersionInfo = 0
    version_infos_schema = 0
    DlInit = 0
    dl_inits_schema = 0
    StatusData = 0
    status_datas_schema = 0
    DlInitInfo = 0
    dl_init_infos_schema = 0
 
    Error = init_global_file_data(module_dir)
    if Error == True:
        logging_text = ' Error Init Global File Data'
        write_log_api_file(module_dir, module, logging_text)
        print(module+logging_text)
    else:
        logging_text = 'File system will be used instead of database'
        write_log_api_file(module_dir, module, logging_text)
        print(module+logging_text)
        logging_text = 'AI Init Values: '+ai_type+' '+ai_type_dl+' '+ai_type_knn+' '+db_type
        write_log_api_file(module_dir, module, logging_text)
        print(module, logging_text)
else:
    pg_port = '5435'        
    Base, session, init_data = db_session(module_dir, db_type, pg_port)
 
    VersionInfo = init_data[0]
    version_infos_schema = init_data[1]
    DlInit = init_data[8]
    dl_inits_schema = init_data[9]
    StatusData = init_data[10]
    status_datas_schema = init_data[11]
    DlInitInfo = init_data[16]
    dl_init_infos_schema = init_data[17]
 
    logging_text = 'AI Init Values: '+ai_type+' '+ai_type_dl+' '+ai_type_knn+' '+db_type
    write_log_api(module_dir, session, StatusData, status_datas_schema, module, logging_text)
 
if __name__ == '__main__':
    #ac_api_rest.run(host="localhost", port=6350, debug=True)                                                          # for insecure communication
    ac_api_rest.run(host="localhost", port=6350, debug=True, ssl_context=('/home/ac/ai_ac_service/api/server_local.crt', '/home/ac/ai_ac_service/api/server_local.key'))     # for secure communication













